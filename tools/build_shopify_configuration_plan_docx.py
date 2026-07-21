from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documentacion" / "PLAN_IMPLEMENTACION_CONFIGURACION_SHOPIFY.md"
OUTPUT = ROOT / "documentacion" / "PLAN_IMPLEMENTACION_CONFIGURACION_SHOPIFY.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "2D6A4F"
GOLD = "7A5A00"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, size=None, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color="C8D0DA", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def apply_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.extend([r_fonts, sz, color])
    r.append(r_pr)
    r.extend([fld_char1, instr_text, fld_char2])
    paragraph._p.append(r)


def add_hyperlink(paragraph, url: str, text: str | None = None):
    text = text or url
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text: str, *, size=11, color=INK):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|<https?://[^>]+>|https?://\S+)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(size - 0.5, 8), color=DARK_BLUE, name="Consolas")
        else:
            url = token[1:-1] if token.startswith("<") else token.rstrip(".,)")
            add_hyperlink(paragraph, url)
            suffix = token[len(url) :]
            if suffix and not token.startswith("<"):
                run = paragraph.add_run(suffix)
                set_run_font(run, size=size, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def set_paragraph_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_custom_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        if kind == "decimal":
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        else:
            lvl_text.set(qn("w:val"), "•" if level == 0 else "–")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), "Calibri")
            r_fonts.set(qn("w:hAnsi"), "Calibri")
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in settings.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("LOVLORY  |  PLAN DE CONFIGURACIÓN SHOPIFY")
    set_run_font(left, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    add_page_number(p)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PLAN OPERATIVO DE IMPLEMENTACIÓN")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Configuración Shopify\nLovLory")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Lista maestra para decidir, configurar, verificar y lanzar la tienda paso a paso")
    set_run_font(r, size=13.5, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    apply_table_geometry(table, [2520, 6840])
    set_table_borders(table, color="D5DCE5")
    values = [
        ("Versión", "1.0 — 21 de julio de 2026"),
        ("Ámbito", "Moneda, mercados, impuestos, pagos, inventario, transporte, checkout, cumplimiento, catálogo y lanzamiento"),
        ("Estado inicial", "Pendiente de completar decisiones previas"),
    ]
    for idx, (label, value) in enumerate(values):
        set_cell_shading(table.cell(idx, 0), LIGHT_BLUE)
        p1 = table.cell(idx, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=10, bold=True, color=DARK_BLUE)
        p2 = table.cell(idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_rich_text(p2, value, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("OBJETIVO")
    set_run_font(r, size=10, bold=True, color=BLUE)

    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="D5DCE5")
    set_cell_shading(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_rich_text(
        p,
        "Convertir la configuración general de Shopify en un proceso controlado y verificable. Cada tarea debe quedar marcada, asignada y respaldada por una evidencia antes de considerarse terminada.",
        size=11,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Orden recomendado")
    set_run_font(r, size=12, bold=True, color=DARK_BLUE)
    for text in (
        "Completar decisiones previas y datos de empresa.",
        "Cerrar moneda, mercados, fiscalidad, pagos y logística.",
        "Aplicar catálogo, políticas, comunicaciones y medición.",
        "Superar las pruebas de aceptación antes de abrir la tienda.",
    ):
        p = doc.add_paragraph()
        set_paragraph_numbering(p, doc._decimal_num_id)
        add_rich_text(p, text, size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Documento vivo · Mantener el Markdown como fuente de seguimiento y el Word como versión compartible.")
    set_run_font(r, size=9.5, italic=True, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def choose_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [1350, 4050, 3960]
    if count == 4:
        return [1350, 2850, 1350, 3810]
    if count == 5:
        return [1300, 2300, 1500, 1300, 2960]
    if count == 6:
        return [900, 900, 2200, 1900, 1400, 2060]
    return [TABLE_WIDTH_DXA // count] * count


def add_wide_form(doc: Document, headers: list[str], rows: list[list[str]]):
    for row in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(row[0] if row else "Registro")
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
        pairs = list(zip(headers[1:], row[1:]))
        table = doc.add_table(rows=(len(pairs) + 1) // 2, cols=4)
        apply_table_geometry(table, [1450, 3230, 1450, 3230])
        set_table_borders(table, color="D5DCE5")
        for idx, (label, value) in enumerate(pairs):
            rr, cc = divmod(idx, 2)
            cc *= 2
            set_cell_shading(table.cell(rr, cc), LIGHT_GRAY)
            p1 = table.cell(rr, cc).paragraphs[0]
            p1.paragraph_format.space_after = Pt(0)
            add_rich_text(p1, label, size=9.5, color=DARK_BLUE)
            p1.runs[0].bold = True
            p2 = table.cell(rr, cc + 1).paragraphs[0]
            p2.paragraph_format.space_after = Pt(0)
            add_rich_text(p2, value, size=9.5)


def add_markdown_table(doc: Document, lines: list[str]):
    parsed = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        parsed.append(cells)
    headers = parsed[0]
    rows = parsed[2:]
    if len(headers) >= 7:
        add_wide_form(doc, headers, rows)
        return

    table = doc.add_table(rows=1, cols=len(headers))
    widths = choose_widths(headers)
    apply_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_rich_text(p, header, size=9.5, color=DARK_BLUE)
        for run in p.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 0 and len(value) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, value, size=9.2)
            if value in {"Pendiente", "En curso", "Terminada", "Aprobado", "Abierta"}:
                for run in p.runs:
                    run.font.color.rgb = rgb(GOLD if value in {"Pendiente", "En curso"} else GREEN)
                    run.bold = True
        apply_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="D5DCE5")
    set_cell_shading(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_rich_text(p, text, size=10.5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def render_markdown(doc: Document, markdown: str):
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## Cómo utilizar"))
    i = start
    current_decimal_num_id = doc._decimal_num_id
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("> "):
            callout_lines = [line[2:]]
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                callout_lines.append(lines[i][2:])
                i += 1
            add_callout(doc, " ".join(callout_lines))
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_rich_text(p, line[4:], size=13, color=BLUE)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, line[3:], size=16, color=BLUE)
            if line.startswith("## Datos pendientes"):
                current_decimal_num_id = doc._pending_num_id
        elif re.match(r"^\s*- \[[ x~!]\] ", line):
            stripped = line.lstrip()
            level = 1 if len(line) - len(stripped) >= 2 else 0
            state = stripped[3]
            text = stripped[6:]
            symbol = {" ": "☐", "x": "☒", "~": "◐", "!": "⚠"}[state]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            r = p.add_run(symbol + " ")
            set_run_font(r, size=11, bold=True, color=GREEN if state == "x" else GOLD if state in {"~", "!"} else DARK_BLUE, name="Segoe UI Symbol")
            add_rich_text(p, text, size=10.5)
        elif re.match(r"^\s*- ", line):
            stripped = line.lstrip()
            level = 1 if len(line) - len(stripped) >= 2 else 0
            p = doc.add_paragraph()
            set_paragraph_numbering(p, doc._bullet_num_id, level)
            add_rich_text(p, stripped[2:], size=10.5)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            set_paragraph_numbering(p, current_decimal_num_id)
            add_rich_text(p, re.sub(r"^\d+\. ", "", line), size=10.5)
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line, size=10.5)
        i += 1


def build():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)
    doc._bullet_num_id = add_custom_numbering(doc, "bullet")
    doc._decimal_num_id = add_custom_numbering(doc, "decimal")
    doc._pending_num_id = add_custom_numbering(doc, "decimal")

    properties = doc.core_properties
    properties.title = "Plan maestro de configuración Shopify — LovLory"
    properties.subject = "Lista de control operativa para configurar y lanzar Shopify"
    properties.author = "LovLory"
    properties.keywords = "Shopify, LovLory, configuración, impuestos, pagos, transporte, lanzamiento"
    properties.comments = "Documento operativo generado a partir del plan de migración LovLory."

    add_cover(doc)
    render_markdown(doc, markdown)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
